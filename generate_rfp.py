"""
AI-Powered Temp Labor RFP Generator
====================================
Takes RFI insights about a client and uses Claude to generate a complete RFP,
then produces an Excel file matching the template exactly.

Usage:
    python generate_rfp.py <rfi_insights.json>

The JSON should contain client context / RFI insights. The AI will generate
all RFP sections (requirements, questionnaire, SLAs, bid sheet roles, mark-ups, etc.)
and output a formatted Excel identical to the reference template.
"""

import json
import sys
import os
import shutil
from copy import copy
from datetime import datetime
from openpyxl import load_workbook
from openpyxl.cell.cell import MergedCell
import anthropic

TEMPLATE_PATH = r"C:\Users\igarg015\Downloads\Shared_Viatris_ Temporary Labor_Agency_RFX (1).xlsx"
OUTPUT_DIR = r"c:\Isha\Claude\Temp labor\output"

# ============================================================
# AI GENERATION - Claude generates the full RFP content
# ============================================================

SYSTEM_PROMPT = """You are a Senior Procurement RFP Specialist with 15+ years of experience in
contingent workforce / temp labor strategy. You generate complete, detailed RFP content
for temporary labor programs.

Given RFI insights about a client, you produce ALL sections of a temp labor RFP as structured JSON.

Your output must be a single valid JSON object with these keys:
{
  "cover_letter": {
    "company_description": "...",
    "objective": "...",
    "scope": "...",
    "confidentiality": "...",
    "evaluation_criteria": "...",
    "rules_of_engagement": "...",
    "bid_instructions": "..."
  },
  "requirements": [
    {"category": "All|Operations|Sterile|...", "subcategory": "...", "text": "..."}
  ],
  "questions": [
    {"category": "Account Management|Compliance & Risk|...", "text": "..."}
  ],
  "slas": [
    {"kpi": "...", "definition": "...", "calculation": "...", "target": "...", "period": "Monthly|Quarterly", "source": "..."}
  ],
  "roles": [
    {"rfp_id": "XX0001", "region": "APAC|EMEA|LATAM|NA", "job_family": "...", "job_title": "...", "description": "...", "country": "...", "site": "...", "demand": 5, "pay_rate": "25.00"}
  ],
  "markups": [
    {"region": "...", "location": "...", "country": "...", "job_family": "...", "demand": 10}
  ],
  "conversions": [
    {"location": "...", "country": "...", "currency": "...", "rate": "0.0125"}
  ]
}

RULES:
1. Generate 40-80 requirements covering: Workforce Readiness & Onboarding, Workforce Management & Staffing Controls, Service Model & General Compliance, Service Delivery, Risk/Incident Management, Reporting & Performance, Health/Safety, Commercial Management, Change/Transition, and any industry-specific needs.

2. Generate 30-53 questionnaire questions across: Account Management, Compliance & Risk, Fulfillment Capability, Geographic Coverage, Pricing & Commercials, Recruitment & Quality, Service Delivery & SLAs, Technology & Reporting.

3. Generate 20-31 SLAs covering: Fill Rate, Time to Fill, Bill Rate Compliance, Invoice Accuracy, On-boarding time, Day-1 No-Show Rate, Early End Turnover, Replacement Turnaround, Quality metrics, Customer Support response, Retention, and operational KPIs.

4. For roles: Generate based on the client's stated needs — job families, locations, pay rates appropriate for the market. Use realistic USD/hr pay rates for each country/role.

5. For mark-ups: One row per unique region + location + job family combination.

6. For conversion rates: Include all non-USD countries in scope.

7. Every requirement must be SPECIFIC and MEASURABLE — no vague language.

8. SLA targets should be industry-standard for temp labor (e.g., Fill Rate ≥90%, Time to Fill ≤14 days, No-Show ≤3%).

9. All content must be tailored to the specific client's industry, locations, and needs.

Output ONLY the JSON — no markdown fences, no explanation."""


def generate_rfp_content(rfi_insights: dict) -> dict:
    """Use Claude to generate complete RFP content from RFI insights."""

    client = anthropic.Anthropic()

    # Build the user message from RFI insights
    user_msg = f"""Generate a complete Temporary Labor RFP for the following client based on these RFI insights:

CLIENT INFORMATION:
- Company: {rfi_insights.get('company', 'Unknown')}
- Industry: {rfi_insights.get('industry', 'General')}
- Company Size: {rfi_insights.get('company_size', 'Not specified')}
- Headquarters: {rfi_insights.get('headquarters', 'Not specified')}

RFP SCOPE:
- Regions: {rfi_insights.get('regions', 'Not specified')}
- Countries: {rfi_insights.get('countries', 'Not specified')}
- Business Units: {rfi_insights.get('business_units', 'Not specified')}
- Total Headcount Needed: {rfi_insights.get('headcount', 'Not specified')}
- Contract Duration: {rfi_insights.get('duration', '2-3 years')}
- Worker Classification: {rfi_insights.get('worker_classification', 'W-2 through staffing agency')}

KEY REQUIREMENTS / RFI INSIGHTS:
{rfi_insights.get('insights', 'No specific insights provided')}

ROLES NEEDED:
{rfi_insights.get('roles_needed', 'Not specified')}

SPECIAL CONSIDERATIONS:
- Industry-specific needs: {rfi_insights.get('industry_specifics', 'None specified')}
- Technology requirements: {rfi_insights.get('tech_requirements', 'VMS to be determined')}
- Compliance focus: {rfi_insights.get('compliance_focus', 'Standard labor law compliance')}
- Supplier model: {rfi_insights.get('supplier_model', 'Multi-supplier')}

TIMELINE:
- RFP Release: {rfi_insights.get('timeline', {}).get('release', 'TBD')}
- Submission Deadline: {rfi_insights.get('timeline', {}).get('deadline', 'TBD')}

COMMERCIAL PREFERENCES:
- Payment Terms: {rfi_insights.get('payment_terms', 'Net 60')}
- Mark-up expectations: {rfi_insights.get('markup_expectations', 'Market competitive')}
- Conversion fee expectation: {rfi_insights.get('conversion_fee', '$0')}

ADDITIONAL CONTEXT:
{rfi_insights.get('additional_context', 'None')}

Generate the complete RFP content as JSON following the structure specified in your instructions. Make it comprehensive, specific to this client's industry and needs, and ready for supplier distribution."""

    print("Calling Claude API to generate RFP content...")
    print(f"  Client: {rfi_insights.get('company', 'Unknown')}")
    print(f"  Industry: {rfi_insights.get('industry', 'General')}")
    print(f"  Regions: {rfi_insights.get('regions', 'Not specified')}")
    print()

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=16000,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_msg}]
    )

    raw = response.content[0].text.strip()

    # Clean up potential markdown fences
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]
    if raw.endswith("```"):
        raw = raw.rsplit("```", 1)[0]
    if raw.startswith("json"):
        raw = raw[4:]

    rfp_content = json.loads(raw)

    print(f"  Generated: {len(rfp_content.get('requirements', []))} requirements")
    print(f"  Generated: {len(rfp_content.get('questions', []))} questions")
    print(f"  Generated: {len(rfp_content.get('slas', []))} SLAs")
    print(f"  Generated: {len(rfp_content.get('roles', []))} roles")
    print(f"  Generated: {len(rfp_content.get('markups', []))} mark-up rows")
    print(f"  Generated: {len(rfp_content.get('conversions', []))} conversion rates")
    print()

    return rfp_content


# ============================================================
# EXCEL GENERATION - Fill template with AI-generated content
# ============================================================

def fill_template(rfp_content: dict, rfi_insights: dict, output_path: str):
    """Copy template and fill with generated content, preserving all formatting."""

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    shutil.copy2(TEMPLATE_PATH, output_path)
    wb = load_workbook(output_path)

    setup = rfi_insights
    cover = rfp_content.get('cover_letter', {})

    # --- TAB 1: Cover Letter ---
    ws = wb['1. Cover Letter']
    # Update title
    title = setup.get('title', 'Temporary Labor Global Agency RFP')
    for row in ws.iter_rows(min_row=1, max_row=5, min_col=2, max_col=3):
        for cell in row:
            if cell.value and 'Temporary Labor' in str(cell.value) and 'RFP' in str(cell.value):
                cell.value = f"{setup.get('company', '')} {title}"

    # Update About section
    if cover.get('company_description'):
        ws['C4'] = cover['company_description']

    # Update objective
    if cover.get('objective'):
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=2, max_col=3):
            for cell in row:
                if cell.value and 'RFP Objective' in str(cell.value):
                    ws.cell(row=cell.row, column=3).value = cover['objective']
                    break

    # Update scope
    if cover.get('scope'):
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=2, max_col=3):
            for cell in row:
                if cell.value and 'RFP Scope' in str(cell.value):
                    ws.cell(row=cell.row, column=3).value = cover['scope']
                    break

    # Update contact
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, min_col=2, max_col=3):
        for cell in row:
            if cell.value and isinstance(cell.value, str) and 'Name:' in cell.value:
                cell.value = f"Name: {setup.get('contact_name', '')}"
            if cell.value and isinstance(cell.value, str) and 'Email:' in cell.value:
                cell.value = f"Email: {setup.get('contact_email', '')}"

    # --- TAB 3: Requirements ---
    ws = wb['3. Requirements']
    requirements = rfp_content.get('requirements', [])
    if requirements:
        data_start_row = None
        for row_idx in range(1, ws.max_row + 1):
            cell_val = ws.cell(row=row_idx, column=2).value
            if cell_val and str(cell_val).strip() == '#':
                data_start_row = row_idx + 1
                break

        if data_start_row:
            # Capture style from first data row
            row_styles = get_row_styles(ws, data_start_row, 8)

            # Clear existing
            for row_idx in range(data_start_row, ws.max_row + 1):
                for col in range(1, 9):
                    safe_clear(ws, row_idx, col)

            # Write new
            for i, req in enumerate(requirements):
                row = data_start_row + i
                safe_write(ws, row, 2, i + 1)
                safe_write(ws, row, 3, req.get('category', 'All'))
                safe_write(ws, row, 4, req.get('subcategory', ''))
                safe_write(ws, row, 5, req.get('text', ''))
                safe_write(ws, row, 6, '')  # Response
                safe_write(ws, row, 7, '')  # Comment
                apply_row_styles(ws, row, row_styles, 8)

    # --- TAB 4: Questionnaire ---
    ws = wb['4. Questionnaire']
    questions = rfp_content.get('questions', [])
    if questions:
        data_start_row = None
        for row_idx in range(1, ws.max_row + 1):
            cell_val = ws.cell(row=row_idx, column=2).value
            if cell_val and str(cell_val).strip() == '#':
                data_start_row = row_idx + 1
                break

        if data_start_row:
            row_styles = get_row_styles(ws, data_start_row, 6)
            for row_idx in range(data_start_row, ws.max_row + 1):
                for col in range(1, 7):
                    safe_clear(ws, row_idx, col)

            for i, q in enumerate(questions):
                row = data_start_row + i
                safe_write(ws, row, 2, i + 1)
                safe_write(ws, row, 3, q.get('category', 'General'))
                safe_write(ws, row, 4, q.get('text', ''))
                safe_write(ws, row, 5, '')  # Response
                apply_row_styles(ws, row, row_styles, 6)

    # --- TAB 5: SLAs ---
    ws = wb['5. Service Level Agreements']
    slas = rfp_content.get('slas', [])
    if slas:
        # Unmerge all to avoid MergedCell errors
        for merge_range in list(ws.merged_cells.ranges):
            ws.unmerge_cells(str(merge_range))

        data_start_row = None
        for row_idx in range(1, ws.max_row + 1):
            cell_val = ws.cell(row=row_idx, column=2).value
            if cell_val and str(cell_val).strip() == 'KPI':
                data_start_row = row_idx + 1
                break

        if data_start_row:
            for row_idx in range(data_start_row, ws.max_row + 1):
                for col in range(1, 10):
                    ws.cell(row=row_idx, column=col).value = None

            for i, sla in enumerate(slas):
                row = data_start_row + i
                ws.cell(row=row, column=2).value = sla.get('kpi', '')
                ws.cell(row=row, column=3).value = sla.get('definition', '')
                ws.cell(row=row, column=4).value = sla.get('calculation', '')
                ws.cell(row=row, column=5).value = sla.get('target', '')
                ws.cell(row=row, column=6).value = sla.get('period', 'Monthly')
                ws.cell(row=row, column=7).value = sla.get('source', '')
                ws.cell(row=row, column=8).value = ''
                ws.cell(row=row, column=9).value = ''

    # --- TABS 6a-6d: Bid Sheets ---
    region_sheets = {
        'APAC': '6.a Bid Sheet_APAC',
        'EMEA': '6.b Bid Sheet_EMEA',
        'LATAM': '6.c Bid Sheet_LATAM',
        'NA': '6.d Bid Sheet_NA'
    }
    roles = rfp_content.get('roles', [])

    for region, sheet_name in region_sheets.items():
        ws = wb[sheet_name]
        region_roles = [r for r in roles if r.get('region') == region]

        data_start_row = None
        for row_idx in range(1, 20):
            cell_val = ws.cell(row=row_idx, column=2).value
            if cell_val and 'RFP ID' in str(cell_val):
                data_start_row = row_idx + 2
                break
        if not data_start_row:
            data_start_row = 15

        row_styles = get_row_styles(ws, data_start_row - 1, 14)

        # Clear existing
        for row_idx in range(data_start_row, ws.max_row + 1):
            for col in range(1, 14):
                safe_clear(ws, row_idx, col)

        # Write roles
        for i, role in enumerate(region_roles):
            row = data_start_row + i
            safe_write(ws, row, 2, role.get('rfp_id', f'{region[0]}C{str(i+1).zfill(4)}'))
            safe_write(ws, row, 3, role.get('job_family', ''))
            safe_write(ws, row, 4, role.get('job_title', ''))
            safe_write(ws, row, 5, role.get('description', ''))
            safe_write(ws, row, 6, role.get('country', ''))
            safe_write(ws, row, 7, role.get('site', ''))
            safe_write(ws, row, 8, role.get('demand', ''))
            safe_write(ws, row, 9, role.get('pay_rate', ''))
            safe_write(ws, row, 10, '')  # Proposed rate
            safe_write(ws, row, 11, '')  # OT multiplier
            safe_write(ws, row, 12, '')  # Bidding?
            safe_write(ws, row, 13, '')  # Comments
            apply_row_styles(ws, row, row_styles, 14)

    # --- TAB 7: Mark-Up ---
    ws = wb['7. Bid Sheet_Mark-Up']
    markups = rfp_content.get('markups', [])
    if markups:
        data_start_row = None
        for row_idx in range(1, 20):
            cell_val = ws.cell(row=row_idx, column=2).value
            if cell_val and 'Region' in str(cell_val):
                data_start_row = row_idx + 2
                break

        if data_start_row:
            row_styles = get_row_styles(ws, data_start_row - 1, 10)
            for row_idx in range(data_start_row, ws.max_row + 1):
                for col in range(1, 10):
                    safe_clear(ws, row_idx, col)

            for i, mu in enumerate(markups):
                row = data_start_row + i
                safe_write(ws, row, 2, mu.get('region', ''))
                safe_write(ws, row, 3, mu.get('location', ''))
                safe_write(ws, row, 4, mu.get('country', ''))
                safe_write(ws, row, 5, mu.get('job_family', ''))
                safe_write(ws, row, 6, mu.get('demand', ''))
                safe_write(ws, row, 7, '')  # Standard %
                safe_write(ws, row, 8, '')  # OT %
                safe_write(ws, row, 9, '')  # Comments
                apply_row_styles(ws, row, row_styles, 10)

    # --- TAB 9: Conversion Rates ---
    ws = wb['9. Conversion Rates']
    conversions = rfp_content.get('conversions', [])
    if conversions:
        data_start_row = 6
        for row_idx in range(data_start_row, ws.max_row + 1):
            for col in range(1, 7):
                safe_clear(ws, row_idx, col)

        for i, conv in enumerate(conversions):
            row = data_start_row + i
            safe_write(ws, row, 3, conv.get('location', ''))
            safe_write(ws, row, 4, conv.get('country', ''))
            safe_write(ws, row, 5, conv.get('currency', ''))
            try:
                safe_write(ws, row, 6, float(conv.get('rate', 0)))
            except (ValueError, TypeError):
                safe_write(ws, row, 6, conv.get('rate', ''))

    # Save
    wb.save(output_path)
    return output_path


# ============================================================
# HELPERS
# ============================================================

def safe_write(ws, row, col, value):
    cell = ws.cell(row=row, column=col)
    if isinstance(cell, MergedCell):
        for merge_range in list(ws.merged_cells.ranges):
            if cell.coordinate in merge_range:
                ws.unmerge_cells(str(merge_range))
                break
        cell = ws.cell(row=row, column=col)
    cell.value = value
    return cell


def safe_clear(ws, row, col):
    cell = ws.cell(row=row, column=col)
    if isinstance(cell, MergedCell):
        return
    cell.value = None


def get_row_styles(ws, row, max_col):
    styles = {}
    for col in range(1, max_col):
        cell = ws.cell(row=row, column=col)
        if not isinstance(cell, MergedCell):
            styles[col] = {
                'font': copy(cell.font),
                'fill': copy(cell.fill),
                'border': copy(cell.border),
                'alignment': copy(cell.alignment),
            }
    return styles


def apply_row_styles(ws, row, styles, max_col):
    for col in range(1, max_col):
        cell = ws.cell(row=row, column=col)
        if isinstance(cell, MergedCell):
            continue
        if col in styles:
            cell.font = copy(styles[col]['font'])
            cell.fill = copy(styles[col]['fill'])
            cell.border = copy(styles[col]['border'])
            cell.alignment = copy(styles[col]['alignment'])


# ============================================================
# MAIN
# ============================================================

def read_rfi_file(filepath: str) -> dict:
    """Read any supported RFI file and return structured insights dict."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.json':
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)

    elif ext == '.txt' or ext == '.csv':
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        return {"company": "Client", "insights": content, "roles_needed": "", "regions": "", "countries": ""}

    elif ext == '.xlsx' or ext == '.xls':
        wb = load_workbook(filepath, data_only=True)
        text_parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                vals = [str(c).strip() for c in row if c]
                if vals:
                    text_parts.append(' | '.join(vals))
        content = '\n'.join(text_parts)
        return {"company": "Client", "insights": content, "roles_needed": "", "regions": "", "countries": ""}

    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            content = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            print("Installing python-docx...")
            os.system("pip install python-docx --quiet")
            from docx import Document
            doc = Document(filepath)
            content = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        return {"company": "Client", "insights": content, "roles_needed": "", "regions": "", "countries": ""}

    elif ext == '.pdf':
        try:
            import PyPDF2
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = '\n'.join([page.extract_text() or '' for page in reader.pages])
        except ImportError:
            print("Installing PyPDF2...")
            os.system("pip install PyPDF2 --quiet")
            import PyPDF2
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = '\n'.join([page.extract_text() or '' for page in reader.pages])
        return {"company": "Client", "insights": content, "roles_needed": "", "regions": "", "countries": ""}

    else:
        # Try reading as text
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return {"company": "Client", "insights": content, "roles_needed": "", "regions": "", "countries": ""}


def main():
    if len(sys.argv) < 2:
        print("=" * 60)
        print("TEMP LABOR RFP GENERATOR")
        print("=" * 60)
        print()
        print("Usage: python generate_rfp.py <rfi_file>")
        print()
        print("Supported file types:")
        print("  .json  - Structured JSON with client fields")
        print("  .txt   - Plain text RFI insights")
        print("  .pdf   - PDF document")
        print("  .docx  - Word document")
        print("  .xlsx  - Excel file")
        print("  .csv   - CSV file")
        print()
        print("The AI will read the file, generate a complete RFP,")
        print("and output an Excel matching the template exactly.")
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"ERROR: File not found: {filepath}")
        sys.exit(1)

    print("=" * 60)
    print("TEMP LABOR RFP GENERATOR")
    print("=" * 60)
    print()
    print(f"Reading RFI file: {filepath}")
    print()

    # Step 1: Read the RFI file (any format)
    rfi_insights = read_rfi_file(filepath)

    # Step 2: AI generates RFP content
    rfp_content = generate_rfp_content(rfi_insights)

    # Step 3: Determine output path
    company = rfi_insights.get('company', 'Client').replace(' ', '_')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    output_path = os.path.join(OUTPUT_DIR, f"{company}_Temp_Labor_RFP_{timestamp}.xlsx")

    # Save intermediate JSON for reference
    intermediate_path = os.path.join(OUTPUT_DIR, f"{company}_rfp_content_{timestamp}.json")
    with open(intermediate_path, 'w', encoding='utf-8') as f:
        json.dump(rfp_content, f, indent=2, ensure_ascii=False)
    print(f"  AI content saved: {intermediate_path}")
    print()

    # Step 4: Fill template → Excel output
    print("Generating Excel from template...")
    fill_template(rfp_content, rfi_insights, output_path)

    print()
    print("=" * 60)
    print(f"DONE! RFP Excel generated:")
    print(f"  {output_path}")
    print(f"  Size: {os.path.getsize(output_path):,} bytes")
    print("=" * 60)

    # Open the file automatically on Windows
    os.startfile(output_path)

    return output_path


if __name__ == '__main__':
    main()
